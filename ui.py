import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QPushButton, QTextEdit, QLabel,
    QVBoxLayout, QWidget, QHBoxLayout, QComboBox, QMessageBox,
    QFileDialog, QSplashScreen, QFrame, QStackedLayout
)
from PyQt5.QtGui import QPixmap, QFont, QIcon
from PyQt5.QtCore import QThread, pyqtSignal, pyqtSlot, QTimer, Qt

# ==== 模块导入 ====
from voice import VoiceInputThread
from gesture import HandGestureThread
from translate import baidu_translate
from ai_writer_thread import AIWriterThread
from docx import Document

class TranslateThread(QThread):
    result_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, text, target_lang, parent=None):
        super().__init__(parent)
        self.text = text
        self.target_lang = target_lang

    def run(self):
        try:
            result = baidu_translate(self.text, from_lang='auto', to_lang=self.target_lang)
            self.result_signal.emit(result)
        except Exception as e:
            self.error_signal.emit(str(e))

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI 辅助写作与翻译平台")
        self.setGeometry(100, 100, 1100, 750)
        self.setWindowIcon(QIcon("app_icon.png"))

        self.setWindowIcon(QIcon("icons/app_icon.ico"))

        self.init_ui()
        self.setStyleSheet(self.load_styles())

    def init_ui(self):
        # ===== 主输入区 =====
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("请输入您的写作内容或翻译文本...")

        # ===== 输出区（默认隐藏） =====
        self.text_output = QTextEdit()
        self.text_output.setReadOnly(True)
        self.text_output.setVisible(False)

        # 展开/折叠输出按钮
        self.expand_btn = QPushButton("🔽 显示输出")
        self.expand_btn.clicked.connect(self.toggle_output_area)

        # ===== 控件组 =====
        self.language_selector = QComboBox()
        self.language_selector.addItems([
            "中文（zh）", "英文（en）", "日文（jp）", "法语（fra）",
            "德语（de）", "俄语（ru）", "韩语（kor）"
        ])

        self.voice_btn = QPushButton("🎤 语音输入")
        self.translate_btn = QPushButton("🌐 翻译")
        self.ai_write_btn = QPushButton("✍️ AI 写作")
        self.gesture_btn = QPushButton("🖐️ 手势识别")
        self.save_btn = QPushButton("💾 保存 Word")

        # ===== 状态提示 =====
        self.status_label = QLabel("状态：准备就绪")
        self.status_label.setAlignment(Qt.AlignRight)
        self.status_label.setObjectName("status_label")

        # ===== 操作栏布局 =====
        controls_layout = QHBoxLayout()
        controls_layout.addWidget(QLabel("目标语言:"))
        controls_layout.addWidget(self.language_selector)
        controls_layout.addStretch()
        controls_layout.addWidget(self.voice_btn)
        controls_layout.addWidget(self.translate_btn)
        controls_layout.addWidget(self.ai_write_btn)
        controls_layout.addWidget(self.gesture_btn)
        controls_layout.addWidget(self.save_btn)

        # ===== 主布局 =====
        layout = QVBoxLayout()
        layout.addWidget(QLabel("✍️ 输入区"))
        layout.addWidget(self.text_input, 4)
        layout.addLayout(controls_layout)
        layout.addWidget(self.expand_btn)
        layout.addWidget(QLabel("📤 输出区"))
        layout.addWidget(self.text_output, 2)
        layout.addWidget(self.status_label)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        # ==== 信号连接 ====
        self.voice_btn.clicked.connect(self.toggle_voice_input)
        self.translate_btn.clicked.connect(self.translate_text)
        self.ai_write_btn.clicked.connect(self.ai_write)
        self.gesture_btn.clicked.connect(self.launch_gesture_module)
        self.save_btn.clicked.connect(self.save_as_word)

        self.voice_thread = None
        self.gesture_thread = None

    def toggle_output_area(self):
        visible = self.text_output.isVisible()
        self.text_output.setVisible(not visible)
        self.expand_btn.setText("🔼 隐藏输出" if not visible else "🔽 显示输出")

    def load_styles(self):
        return """
        QMainWindow {
            background-color: #f4f7fa;
        }
        QTextEdit {
            border: 1px solid #ccc;
            border-radius: 10px;
            padding: 10px;
            background-color: white;
            font-size: 15px;
        }
        QPushButton {
            background-color: #4A90E2;
            color: white;
            border: none;
            padding: 8px 14px;
            font-size: 14px;
            border-radius: 8px;
        }
        QPushButton:hover {
            background-color: #357ABD;
        }
        QLabel {
            font-weight: bold;
            font-size: 14px;
        }
        QComboBox {
            padding: 6px;
            font-size: 14px;
            border-radius: 6px;
        }
        QLabel#status_label {
            color: #666;
            font-style: italic;
            margin-top: 8px;
        }
        """

    def save_as_word(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "保存为 Word 文件", "", "Word 文件 (*.docx)")
        if file_name:
            doc = Document()
            doc.add_heading("输入内容", level=2)
            doc.add_paragraph(self.text_input.toPlainText())
            doc.add_heading("翻译/AI输出", level=2)
            doc.add_paragraph(self.text_output.toPlainText())
            doc.save(file_name)
            QMessageBox.information(self, "保存成功", f"文件已保存到：{file_name}")

    def toggle_voice_input(self):
        if self.voice_thread is None:
            self.status_label.setText("状态：语音识别中...")
            self.voice_thread = VoiceInputThread()
            self.voice_thread.result_signal.connect(self.on_voice_result)
            self.voice_thread.start()
            self.voice_btn.setText("🛑 停止语音输入")
        else:
            self.voice_thread.stop()
            self.voice_thread = None
            self.voice_btn.setText("🎤 语音输入")
            self.status_label.setText("状态：已停止语音识别")

    @pyqtSlot(str)
    def on_voice_result(self, text):
        current = self.text_input.toPlainText()
        self.text_input.setPlainText(current + "\n" + text)

    def translate_text(self):
        text = self.text_input.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "警告", "请输入要翻译的文本")
            return
        lang = self.language_selector.currentText().split("（")[1].split("）")[0]
        self.status_label.setText("状态：翻译中...")
        self.translate_thread = TranslateThread(text, lang)
        self.translate_thread.result_signal.connect(self.on_translate_result)
        self.translate_thread.error_signal.connect(self.on_translate_error)
        self.translate_thread.start()

    @pyqtSlot(str)
    def on_translate_result(self, result):
        self.text_output.setPlainText(result)
        self.status_label.setText("状态：翻译完成")

    @pyqtSlot(str)
    def on_translate_error(self, error):
        QMessageBox.critical(self, "翻译失败", error)
        self.status_label.setText("状态：翻译出错")

    def ai_write(self):
        text = self.text_input.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "警告", "请输入 AI 写作提示")
            return
        self.status_label.setText("状态：AI 写作中...")
        self.ai_thread = AIWriterThread(text)
        self.ai_thread.result_signal.connect(self.on_ai_result)
        self.ai_thread.error_signal.connect(self.on_ai_error)
        self.ai_thread.start()

    @pyqtSlot(str)
    def on_ai_result(self, result):
        self.text_output.setPlainText(result)
        self.status_label.setText("状态：AI 写作完成")

    @pyqtSlot(str)
    def on_ai_error(self, error):
        QMessageBox.critical(self, "AI 写作失败", error)
        self.status_label.setText("状态：AI 写作出错")

    def launch_gesture_module(self):
        if self.gesture_thread is None:
            self.status_label.setText("状态：手势识别中...")
            self.gesture_thread = HandGestureThread()
            self.gesture_thread.result_signal.connect(self.on_gesture_result)
            self.gesture_thread.start()
            self.gesture_btn.setText("🛑 停止手势识别")
        else:
            self.gesture_thread.stop()
            self.gesture_thread = None
            self.gesture_btn.setText("🖐️ 手势识别")
            self.status_label.setText("状态：手势已停止")

    @pyqtSlot(str)
    def on_gesture_result(self, gesture):
        current = self.text_input.toPlainText()
        self.text_input.setPlainText(current + "\n[手势]：" + gesture)


class SplashScreen(QSplashScreen):
    def __init__(self):
        pixmap = QPixmap("splash_image.png")
        super().__init__(pixmap)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint)
        self.setFont(QFont("微软雅黑", 12))
        self.showMessage("欢迎使用 AI 辅助写作与翻译平台", Qt.AlignBottom | Qt.AlignCenter, Qt.white)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("icons/app_icon.ico"))

    #splash = SplashScreen()
    #splash.show()

    #def show_main():
    #    global win
    #    win = MainWindow()
    #    win.show()
    #    splash.close()

    #    QTimer.singleShot(2500, show_main)
    # 直接进入主界面
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())

