import sys
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QSplitter,
    QListWidget, QListWidgetItem, QTextEdit, QPushButton, QLabel, QComboBox,
    QMenuBar, QMenu, QAction, QToolBar, QStatusBar, QMessageBox, QInputDialog,
    QShortcut, QFileDialog, QSplashScreen, QFrame, QStackedLayout,
    QAbstractItemView
)
from PyQt5.QtGui import QPixmap, QFont, QIcon, QKeySequence
from PyQt5.QtCore import (
    QThread, pyqtSignal, pyqtSlot, QTimer, Qt,
    QEvent, QSize, QPropertyAnimation, QEasingCurve
)


# ==== 模块导入 ====
from voice import VoiceInputThread
from gesture import HandGestureThread
from translate import baidu_translate
from ai_writer_thread import AIWriterThread
from docx import Document

class TranslateThread(QThread):
    result_signal = pyqtSignal(str)
    error_signal = pyqtSignal(str)

    def __init__(self, text, target_lang, parent=None):
        super().__init__(parent)
        self.text = text
        self.target_lang = target_lang

    def run(self):
        try:
            result = baidu_translate(self.text, from_lang='auto', to_lang=self.target_lang)
            self.result_signal.emit(result)
        except Exception as e:
            self.error_signal.emit(str(e))

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("高级交互示例软件")
        self.resize(1100, 700)
        self.setWindowIcon(QIcon("icons/app_icon.ico"))  # 请替换为你自己的图标路径

        self.init_menu_toolbar()
        self.init_ui()
        self.init_status_bar()
        self.apply_dark_theme()

        # 历史记录折叠状态
        self.history_expanded = True

    def init_menu_toolbar(self):
        # 菜单栏
        menubar = self.menuBar()

        file_menu = menubar.addMenu("&文件")
        new_act = QAction("新建", self)
        save_act = QAction("保存", self)
        exit_act = QAction("退出", self)
        exit_act.triggered.connect(self.close)

        file_menu.addAction(new_act)
        file_menu.addAction(save_act)
        file_menu.addSeparator()
        file_menu.addAction(exit_act)

        view_menu = menubar.addMenu("&视图")
        self.toggle_history_act = QAction("显示/隐藏历史记录", self, checkable=True)
        self.toggle_history_act.setChecked(True)
        self.toggle_history_act.triggered.connect(self.toggle_history_panel)
        view_menu.addAction(self.toggle_history_act)

        help_menu = menubar.addMenu("&帮助")
        about_act = QAction("关于", self)
        about_act.triggered.connect(self.show_about_dialog)
        help_menu.addAction(about_act)

        # 工具栏
        toolbar = QToolBar("主工具栏")
        toolbar.setIconSize(QSize(16,16))

        self.addToolBar(toolbar)

        toolbar.addAction(new_act)
        toolbar.addAction(save_act)

        # 示例按钮
        toggle_history_btn = QAction(QIcon("icons/bookmark.png"), "显示/隐藏历史", self)
        toggle_history_btn.triggered.connect(self.toggle_history_panel)
        toolbar.addAction(toggle_history_btn)

    def init_ui(self):
        # 历史记录列表
        self.history_list = QListWidget()
        self.history_list.setSelectionMode(QAbstractItemView.SingleSelection)
        self.history_list.itemDoubleClicked.connect(self.rename_history_item)
        self.history_list.setContextMenuPolicy(Qt.CustomContextMenu)
        self.history_list.customContextMenuRequested.connect(self.history_context_menu)

        history_layout = QVBoxLayout()
        history_layout.setContentsMargins(5,5,5,5)
        history_layout.addWidget(QLabel("历史记录"))
        history_layout.addWidget(self.history_list)
        history_frame = QFrame()
        history_frame.setLayout(history_layout)
        history_frame.setMaximumWidth(300)
        history_frame.setMinimumWidth(0)
        self.history_frame = history_frame

        # 输入输出区域的右半部分
        # 输入框
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("请输入内容，Ctrl+Enter 发送")
        self.text_input.installEventFilter(self)  # 用于捕获快捷键

        # 输出框，富文本只读
        self.text_output = QTextEdit()
        self.text_output.setReadOnly(True)
        self.text_output.setPlaceholderText("输出区域")
        self.text_output.setVisible(False)

        # 语言选择、按钮组
        self.language_selector = QComboBox()
        self.language_selector.addItems(["中文（zh）", "英文（en）", "日文（jp）", "法语（fra）",
                                         "德语（de）", "俄语（ru）", "韩语(kor)"])

        self.voice_btn = QPushButton("语音")
        self.voice_btn.setIcon(QIcon("icons/mic.png"))
        self.translate_btn = QPushButton("翻译")
        self.translate_btn.setIcon(QIcon("icons/translate.png"))
        self.ai_btn = QPushButton("AI写作")
        self.ai_btn.setIcon(QIcon("icons/ai.png"))
        self.gesture_btn = QPushButton("手势")
        self.gesture_btn.setIcon(QIcon("icons/hand.png"))

        # 展开输出区按钮
        self.expand_btn = QPushButton("🔽 显示输出区域")
        self.expand_btn.clicked.connect(self.toggle_output_area)

        # 右侧上方按钮和选择行
        top_button_row = QHBoxLayout()
        self.toggle_history_btn = QPushButton("📚 显示/隐藏历史记录")
        self.toggle_history_btn.clicked.connect(self.toggle_history_panel)
        top_button_row.addWidget(self.toggle_history_btn)
        top_button_row.addStretch()

        # 控件按钮行
        button_row = QHBoxLayout()
        button_row.addWidget(QLabel("目标语言："))
        button_row.addWidget(self.language_selector)
        button_row.addWidget(self.voice_btn)
        button_row.addWidget(self.translate_btn)
        button_row.addWidget(self.ai_btn)
        button_row.addWidget(self.gesture_btn)

        # 右侧主布局垂直方向
        right_layout = QVBoxLayout()
        right_layout.addLayout(top_button_row)
        right_layout.addWidget(QLabel("输入内容："))
        right_layout.addWidget(self.text_input, 4)
        right_layout.addLayout(button_row)
        right_layout.addWidget(self.expand_btn)
        right_layout.addWidget(QLabel("输出内容："))
        right_layout.addWidget(self.text_output, 2)

        right_widget = QWidget()
        right_widget.setLayout(right_layout)

        # 主界面分割器，左右
        self.main_splitter = QSplitter(Qt.Horizontal)
        self.main_splitter.addWidget(self.history_frame)
        self.main_splitter.addWidget(right_widget)
        self.main_splitter.setSizes([280, 800])

        self.setCentralWidget(self.main_splitter)

        # 绑定功能按钮示例（目前只是状态提示）
        self.voice_btn.clicked.connect(lambda: self.update_status("语音功能点击"))
        self.translate_btn.clicked.connect(lambda: self.update_status("翻译功能点击"))
        self.ai_btn.clicked.connect(lambda: self.update_status("AI写作功能点击"))
        self.gesture_btn.clicked.connect(lambda: self.update_status("手势功能点击"))

        # 快捷键 Ctrl+Enter 发送
        self.send_shortcut = QShortcut(QKeySequence("Ctrl+Return"), self)
        self.send_shortcut.activated.connect(self.simulate_send)

    def init_status_bar(self):
        self.statusbar = QStatusBar()
        self.setStatusBar(self.statusbar)
        self.status_label = QLabel("状态：就绪")
        self.statusbar.addWidget(self.status_label)

    def update_status(self, text):
        self.status_label.setText(f"状态：{text}")

    def simulate_send(self):
        content = self.text_input.toPlainText().strip()
        if content:
            # 简单示例，向输出添加富文本，带时间
            from datetime import datetime
            now = datetime.now().strftime("%H:%M:%S")
            self.text_output.setVisible(True)
            self.expand_btn.setText("🔼 隐藏输出区域")

            self.text_output.append(f'<b><font color="green">[{now}] 输入：</font></b>{content}')
            self.text_output.append(f'<b><font color="blue">[{now}] 输出：</font></b>模拟回应：{content[::-1]}')  # 反转字符串做示范
            self.text_input.clear()
            self.update_status("发送成功")
        else:
            self.update_status("输入为空，无法发送")

    def toggle_output_area(self):
        if self.text_output.isVisible():
            self.text_output.setVisible(False)
            self.expand_btn.setText("🔽 显示输出区域")
        else:
            self.text_output.setVisible(True)
            self.expand_btn.setText("🔼 隐藏输出区域")

    def toggle_history_panel(self):
        # 用动画折叠展开历史面板
        if self.history_expanded:
            # 折叠
            self.animate_history_panel(280, 0)
            self.history_expanded = False
            self.toggle_history_btn.setText("📚 显示历史记录")
            self.toggle_history_act.setChecked(False)
        else:
            # 展开
            self.animate_history_panel(0, 280)
            self.history_expanded = True
            self.toggle_history_btn.setText("📚 隐藏历史记录")
            self.toggle_history_act.setChecked(True)

    def animate_history_panel(self, start_width, end_width):
        animation = QPropertyAnimation(self.history_frame, b"maximumWidth")
        animation.setDuration(300)
        animation.setStartValue(start_width)
        animation.setEndValue(end_width)
        animation.setEasingCurve(QEasingCurve.InOutCubic)
        animation.start()
        # 需要保持引用防止GC
        self.animation = animation

    def rename_history_item(self, item):
        old_text = item.text()
        new_text, ok = QInputDialog.getText(self, "重命名历史项", "请输入新名称:", text=old_text)
        if ok and new_text.strip():
            item.setText(new_text.strip())
            self.update_status(f"重命名为: {new_text.strip()}")

    def history_context_menu(self, pos):
        item = self.history_list.itemAt(pos)
        if item is None:
            return
        menu = QMenu()
        rename_act = QAction("重命名", self)
        delete_act = QAction("删除", self)

        rename_act.triggered.connect(lambda: self.rename_history_item(item))
        delete_act.triggered.connect(lambda: self.delete_history_item(item))

        menu.addAction(rename_act)
        menu.addAction(delete_act)
        menu.exec_(self.history_list.viewport().mapToGlobal(pos))

    def delete_history_item(self, item):
        ret = QMessageBox.question(self, "确认删除", f"确认删除历史项 '{item.text()}' ?",
                                   QMessageBox.Yes | QMessageBox.No)
        if ret == QMessageBox.Yes:
            row = self.history_list.row(item)
            self.history_list.takeItem(row)
            self.update_status(f"已删除历史项：{item.text()}")

    def show_about_dialog(self):
        QMessageBox.information(self, "关于", "高级交互示例软件\n由 ChatGPT 助手设计\n2025")

    def eventFilter(self, source, event):
        # Ctrl+Enter 发送事件捕获（也可用快捷键实现，这里做双保险）
        if source == self.text_input and event.type() == QEvent.KeyPress:
            if event.key() == Qt.Key_Return and event.modifiers() == Qt.ControlModifier:
                self.simulate_send()
                return True
        return super().eventFilter(source, event)

    def apply_dark_theme(self):
        dark_style = """
        QMainWindow {
            background-color: #2b2b2b;
            color: #f0f0f0;
        }
        QTextEdit {
            background-color: #3c3f41;
            color: #f0f0f0;
            border: 1px solid #555555;
            border-radius: 4px;
            font-size: 14px;
        }
        QListWidget {
            background-color: #3c3f41;
            color: #f0f0f0;
            border: 1px solid #555555;
            border-radius: 4px;
        }
        QPushButton {
            background-color: #444444;
            color: #f0f0f0;
            border: none;
            padding: 6px 12px;
            border-radius: 4px;
            font-weight: bold;
        }
        QPushButton:hover {
            background-color: #666666;
        }
        QLabel {
            font-weight: bold;
            font-size: 14px;
            color: #dddddd;
        }
        QComboBox {
            background-color: #3c3f41;
            color: #f0f0f0;
            border: 1px solid #555555;
            border-radius: 4px;
            padding: 2px 8px;
        }
        QMenuBar {
            background-color: #2b2b2b;
            color: #f0f0f0;
        }
        QMenuBar::item {
            background-color: #2b2b2b;
            color: #f0f0f0;
            padding: 4px 12px;
        }
        QMenuBar::item:selected {
            background-color: #555555;
        }
        QMenu {
            background-color: #3c3f41;
            color: #f0f0f0;
        }
        QMenu::item:selected {
            background-color: #555555;
        }
        QStatusBar {
            background-color: #2b2b2b;
            color: #f0f0f0;
        }
        """
        self.setStyleSheet(dark_style)

    def save_as_word(self):
        file_name, _ = QFileDialog.getSaveFileName(self, "保存为 Word 文件", "", "Word 文件 (*.docx)")
        if file_name:
            doc = Document()
            doc.add_heading("输入内容", level=2)
            doc.add_paragraph(self.text_input.toPlainText())
            doc.add_heading("输出内容", level=2)
            doc.add_paragraph(self.text_output.toPlainText())
            doc.save(file_name)
            QMessageBox.information(self, "保存成功", f"已保存到：{file_name}")

    def toggle_voice_input(self):
        if self.voice_thread is None:
            self.status_label.setText("状态：语音识别中...")
            self.voice_thread = VoiceInputThread()
            self.voice_thread.result_signal.connect(self.on_voice_result)
            self.voice_thread.start()
            self.voice_btn.setText("🛑 停止语音")
        else:
            self.voice_thread.stop()
            self.voice_thread = None
            self.voice_btn.setText("🎤 语音")
            self.status_label.setText("状态：语音停止")

    def toggle_gesture_input(self):
        if self.gesture_thread is None:
            self.status_label.setText("状态：手势识别中...")
            self.gesture_thread = HandGestureThread()
            self.gesture_thread.result_signal.connect(self.on_gesture_result)
            self.gesture_thread.start()
            self.gesture_btn.setText("🛑 停止手势")
        else:
            self.gesture_thread.stop()
            self.gesture_thread = None
            self.gesture_btn.setText("🖐️ 手势")
            self.status_label.setText("状态：手势识别停止")

    def translate_text(self):
        text = self.text_input.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "警告", "请输入要翻译的文本")
            return
        lang = self.language_selector.currentText().split("（")[1].split("）")[0]
        self.status_label.setText("状态：翻译中...")
        self.translate_thread = TranslateThread(text, lang)
        self.translate_thread.result_signal.connect(self.on_translate_result)
        self.translate_thread.error_signal.connect(self.on_translate_error)
        self.translate_thread.start()

    def ai_write(self):
        text = self.text_input.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "警告", "请输入 AI 写作提示")
            return
        self.status_label.setText("状态：AI 写作中...")
        self.ai_thread = AIWriterThread(text)
        self.ai_thread.result_signal.connect(self.on_ai_result)
        self.ai_thread.error_signal.connect(self.on_ai_error)
        self.ai_thread.start()

    @pyqtSlot(str)
    def on_voice_result(self, result):
        self.text_input.append(result)
        self.add_history_item("语音输入")

    @pyqtSlot(str)
    def on_gesture_result(self, result):
        self.text_input.append("[手势]：" + result)
        self.add_history_item("手势输入")

    @pyqtSlot(str)
    def on_translate_result(self, result):
        self.text_output.setPlainText(result)
        self.status_label.setText("状态：翻译完成")
        self.add_history_item("翻译")

    @pyqtSlot(str)
    def on_translate_error(self, err):
        QMessageBox.critical(self, "翻译失败", err)
        self.status_label.setText("状态：翻译出错")

    @pyqtSlot(str)
    def on_ai_result(self, result):
        self.text_output.setPlainText(result)
        self.status_label.setText("状态：AI 写作完成")
        self.add_history_item("AI 写作")

    @pyqtSlot(str)
    def on_ai_error(self, err):
        QMessageBox.critical(self, "AI 写作失败", err)
        self.status_label.setText("状态：写作出错")

    def add_history_item(self, prefix):
        item = QListWidgetItem(f"{prefix} - 项目 {len(self.history_data) + 1}")
        self.history_data.append(item.text())
        self.history_list.addItem(item)

    def rename_history_item(self, item):
        new_text, ok = QInputDialog.getText(self, "重命名历史记录", "新名称：", text=item.text())
        if ok and new_text.strip():
            item.setText(new_text.strip())


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("app_icon.ico"))
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())
